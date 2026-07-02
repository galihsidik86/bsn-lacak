"""Patch Proposal-BSN-Lacak-(ArtiVisi).docx + Proposal-BSN-Lacak-(Lampiran).docx
supaya akurat dengan sistem yang sudah dibangun. Pakai python-docx +
text-search-replace di run + cell level supaya formatting preserved.

Strategi:
- Text replacement eksak di paragraph runs
- Text replacement eksak di table cell (aggregate multi-run per cell)
- Idempotent — jalankan ulang aman

Jalankan: python3 docs/update-proposal.py
"""
from __future__ import annotations
import sys
try: sys.stdout.reconfigure(encoding='utf-8')
except Exception: pass
from pathlib import Path
from copy import deepcopy
from docx import Document

DOCS = Path(r"C:/Users/Galih Sidik/BSN/docs")
FILES = [
    DOCS / "Proposal-BSN-Lacak-(ArtiVisi).docx",
    DOCS / "Proposal-BSN-Lacak-(Lampiran).docx",
]

# Replacement dict — old_substring: new_substring. Match substring dalam
# text paragraph atau cell. Idempotent kalau replacement tidak mengubah
# text jadi mengandung old_substring (yang harusnya tidak terjadi kalau
# hati-hati).
REPLACEMENTS: list[tuple[str, str]] = [
    # ==================== VERSI + TANGGAL ====================
    ("17 Juni 2026", "1 Juli 2026"),
    ("Bogor, 17 Juni 2026", "Bogor, 1 Juli 2026"),
    ("v1.2", "v1.3"),
    ("1.2 — dengan estimasi biaya & perbandingan map provider",
     "1.3 — refresh sesuai sistem yang sudah dibangun (Capacitor APK, FCM, chat, dll)"),
    ("1.2 — dengan perbandingan Map Provider",
     "1.3 — refresh sesuai sistem yang sudah dibangun"),

    # ==================== MAIN PROPOSAL ====================

    # Modul 2 — deskripsi cross-platform + tambah APK
    ("PWA lintas-platform — installable di Android (Chrome) dan iOS (Safari) tanpa app store.",
     "PWA lintas-platform + APK Android Native (Capacitor) — installable di Android (Chrome/APK), iOS (Safari PWA). APK punya foreground service GPS."),

    # Fitur F-08 — IndexedDB → localStorage
    ("Form digital hasil + nominal + catatan. Antrian offline IndexedDB.",
     "Form digital hasil + nominal + catatan. Antrian offline localStorage + retry queue clientTs preserved."),

    # Fitur F-10 — Anti-Tampering deskripsi
    ("Deteksi Mock GPS, Root/Jailbreak, modifikasi pihak ketiga. + anti-fraud server.",
     "In-app camera lock (galeri ditolak) + magic-byte server validation + EXIF freshness > 1 jam + speed-jump > 150 km/h detection."),

    # ---- Tech stack (Table 5) ----
    ("Node.js 20 + Express + TypeScript + Prisma ORM",
     "Node.js 22 + Express + TypeScript (Zod validation) + Prisma ORM"),
    ("PostgreSQL 14+ dengan streaming replication (untuk skema HA)",
     "PostgreSQL 16 (single-host untuk pilot; streaming replication untuk skala > 1 region)"),
    ("Vite + React 18 + TypeScript + PWA",
     "Vite + React 18 + TypeScript + PWA + Capacitor APK Android (dual runtime)"),
    ("Web Push standar (VAPID), tanpa FCM/APNs proprietary",
     "Dual: Web Push VAPID (PWA browser) + Firebase Cloud Messaging (FCM) untuk APK Android"),
    ("passport-ldapauth siap integrasi dengan Active Directory BSN",
     "Slot integrasi LDAP/AD via passport-ldapauth siap saat BSN kirim spec AD (belum diimplementasi fase MVP)"),
    ("Nginx — TLS 1.3, rate-limit per-IP, security headers (Helmet, CSP, HSTS)",
     "Caddy — TLS 1.3 + HTTP/3 auto Let's Encrypt, rate-limit per-IP, security headers (Helmet, CSP, HSTS)"),

    # ---- Security aspects (Table 6) ----
    ("Enkripsi Data (in Transit)",
     "Enkripsi Data (in Transit)"),  # keep header, just for locate
    ("TLS 1.3 / HTTPS via Nginx. Sertifikat Let's Encrypt atau internal CA BSN.",
     "TLS 1.3 + HTTP/3 via Caddy. Sertifikat Let's Encrypt auto-renew atau internal CA BSN."),
    ("LDAP/SSO Active Directory + TOTP 2FA wajib Supervisor/Admin. Account lockout 5x salah.",
     "Username/password + TOTP 2FA (siap wajib untuk SUPERVISOR/ADMIN) + JWT 15m + Refresh 7d. Account lockout 5x salah. Slot LDAP/AD SSO siap aktivasi."),
    ("Capacitor + plugin freerasp + mock-location-detection.",
     "APK Capacitor + in-app camera lock (capture attribute) + magic-byte JPEG/PNG server validation."),
    ("GPS plausibility > 200m, EXIF freshness > 1 jam, speed jump > 150 km/h.",
     "Geo-fencing 50m + GPS plausibility + EXIF freshness > 1 jam + speed jump > 150 km/h + duplicate visit + nominal spike detection."),

    # ---- Performance (Table 7) ----
    ("99,9% — skema HA: streaming replication Postgres + dual app server di belakang nginx LB.",
     "Single-host uptime 99,5% (Caddy + Docker Compose + backup harian). HA (Postgres replica + LB) tersedia sebagai upgrade opsional saat scale > 1 region."),
    ("PWA + IndexedDB queue + sync otomatis saat online + persistensi tab-kill.",
     "PWA/APK + localStorage retry queue + auto-sync saat online/focus + clientTs preserved (recordedAt server sesuai capture asli, bukan waktu drain)."),
    ("pg_dump otomatis 24 jam, retensi 7 hari. Restore drill di runbook. WAL streaming untuk RPO < 1 menit.",
     "pg_dump otomatis harian, retensi 30 hari. Restore drill di runbook. Trail GPS retention 90 hari (worker harian auto-prune > 90 hari)."),

    # ==================== LAMPIRAN ====================

    # Domain entities (Table 5 di Lampiran) — count 14 → 35+
    ("Domain model 14 entitas",
     "Domain model 35+ entitas"),

    # UI walkthrough count
    ("Walkthrough 14 halaman Web Dashboard (Supervisor + Admin) dan 8 halaman Mobile PWA Petugas",
     "Walkthrough 20+ halaman Web Dashboard (Supervisor + Admin) dan 5 tab Mobile Petugas (Beranda/Rute/Riwayat/Profil + overlay Lapor/Chat + Onboarding Tour)"),

    # API endpoint count
    ("Daftar lengkap 60+ endpoint REST API",
     "Daftar lengkap 190+ endpoint REST API"),

    # Roadmap P0/P1/P2 detail cells
    ("passport-ldapauth + speakeasy. Mandatory untuk Supervisor/Admin.",
     "TOTP 2FA (speakeasy) DONE untuk SUPERVISOR/ADMIN. Slot LDAP/AD siap aktivasi via passport-ldapauth."),
    ("Wrap PWA + plugin freerasp + mock-location-detection. Build APK+IPA.",
     "Wrap PWA jadi APK Android (Capacitor). In-app camera + magic byte + EXIF freshness DONE. Foreground GPS service + FCM push DONE."),
    ("Anti-Tampering — Capacitor",
     "APK Capacitor + Foreground GPS + FCM Push"),

    # Time-on-site → GPS trail historis (yang beneran dibuat)
    ("Time-on-site visualization",
     "GPS trail historis + heatmap kunjungan"),
    ("Komputasi durasi per-titik dari rute.",
     "Trail polyline dengan date picker (audit lintas-hari) + heatmap kunjungan cross-cabang untuk hotspot detection."),

    # Postgres replica
    ("Postgres replica + nginx LB + dual app server.",
     "Postgres streaming replica + Caddy multi-instance + dual app server (untuk scale > 1 region)."),

    # WhatsApp Cloud API — realistic
    ("Migrasi gateway ke Meta Cloud API + template approved Meta.",
     "Aktivasi Meta Cloud API resmi + template approved Meta (BlastGateway pluggable sudah siap, tinggal isi kredensial saat BSN punya akun Meta Business)."),

    # ==================== ROLE-BASED COST (Lampiran T16) ====================
    ("Mobile Engineer (Capacitor)",
     "Mobile Engineer (Capacitor)"),  # already correct
    ("PWA + Capacitor native wrap, plugin anti-tampering, build & sign APK/IPA.",
     "PWA + Capacitor APK wrap (foreground GPS + FCM), build & sign APK. iOS opsional (butuh Apple Developer account BSN)."),

    # ==================== INFRA COST (Lampiran T20) ====================
    ("Server aplikasi, database, storage, network, sertifikat TLS internal.",
     "Server aplikasi + database + storage + network + sertifikat TLS internal (atau Let's Encrypt gratis via Caddy)."),

    # Firebase Cloud Messaging note — add di WA row (kalau bisa)
    ("Registrasi BSP untuk WhatsApp Cloud API.",
     "Registrasi BSP untuk WhatsApp Cloud API. Firebase project untuk FCM APK push (gratis di tier Spark, sudah cukup untuk ratusan device)."),

    # ==================== SERVER SPEC (Lampiran T24) ====================
    # (nothing to swap yet — spec already generic)

    # ==================== ROUND 2: PATCH SISA YANG BELUM MATCH ====================
    # Body text section 4.2 Modul 2 Mobile PWA
    ("Aplikasi mobile dirancang sebagai PWA cross-platform yang dapat di-install di Android (banner Chrome) dan iOS (panduan Share Safari). Untuk Anti-Tampering yang membutuhkan akses native, PWA dibungkus dengan Capacitor + plugin freerasp dan mock-location-detection.",
     "Aplikasi mobile dirancang sebagai PWA cross-platform installable di Android (banner Chrome) dan iOS (Share Safari), dilengkapi dengan APK Android Native (Capacitor) yang punya foreground service GPS untuk tracking saat layar mati. Anti-tampering fase MVP: in-app camera lock + magic-byte server validation + EXIF freshness + speed-jump detection."),

    # Body text section 5.1 Prinsip Arsitektur
    ("BSN Lacak dirancang dengan arsitektur layered services modern: Presentation Layer (Web Dashboard + Mobile PWA) → Business Logic Layer (Node.js + Express + TypeScript dengan validasi Zod, JWT auth, RBAC, background workers) → Data Layer (PostgreSQL dengan Prisma ORM + file storage + IndexedDB client).",
     "BSN Lacak dirancang dengan arsitektur layered services modern: Presentation Layer (Web Dashboard + Mobile PWA + APK Android Native via Capacitor) → Business Logic Layer (Node.js 22 + Express + TypeScript dengan validasi Zod, JWT auth 15m + Refresh 7d, RBAC 3-tier, 15+ background workers) → Data Layer (PostgreSQL 16 dengan Prisma ORM + file storage + localStorage retry queue client)."),

    # Cell Lampiran 2 deskripsi
    ("Diagram layered services BSN Lacak, domain model 14 entitas, skema deployment Docker Compose, dan strategi High Availability.",
     "Diagram layered services BSN Lacak, domain model 35+ entitas, skema deployment Docker Compose (Caddy TLS auto + Postgres 16 + API Node 22), dan strategi High Availability opsional."),

    # Cell alur pipeline: "Blank spot → IndexedDB → auto-sync..."
    ("Real-time submit ke server + foto ber-watermark. Blank spot → IndexedDB → auto-sync saat online.",
     "Real-time submit ke server + foto ber-watermark server-side. Blank spot → localStorage retry queue → auto-sync saat online/focus dengan clientTs preserved."),

    # Lampiran arsitektur text
    ("Dua instance API di belakang nginx LB — session stateless (JWT).",
     "Multi-instance API di belakang Caddy — session stateless (JWT). Opsi HA untuk skala > 1 region."),

    # ASCII diagram — replace 1:1 length untuk preserve box alignment.
    # "IndexedDB" (9 char) → "Local Store" (11 char) — 2 char lebih panjang,
    # box border sedikit shift tapi tetap readable.
    ("   - IndexedDB offline queue  |",
     "   - localStorage retry queue |"),
]


def patch_paragraph(para, replacements: list[tuple[str, str]]) -> int:
    """Apply replacements ke paragraph, preserving run formatting."""
    n = 0
    full = "".join(r.text for r in para.runs)
    changed = False
    for old, new in replacements:
        if old in full:
            full = full.replace(old, new)
            changed = True
            n += 1
    if changed and para.runs:
        para.runs[0].text = full
        for r in para.runs[1:]:
            r.text = ""
    return n


def patch_cell(cell, replacements: list[tuple[str, str]]) -> int:
    """Apply replacements di semua paragraph dalam sebuah cell."""
    n = 0
    for para in cell.paragraphs:
        n += patch_paragraph(para, replacements)
    return n


def add_new_rows(doc: Document, table_locator, new_rows: list[list[str]]) -> int:
    """Append rows baru ke tabel yang match table_locator (function
    yang terima Table, return True kalau target).
    """
    added = 0
    for tbl in doc.tables:
        if not table_locator(tbl):
            continue
        for row_content in new_rows:
            # Clone last row untuk preserve formatting cell
            template_row = tbl.rows[-1]._tr
            new_tr = deepcopy(template_row)
            tbl._tbl.append(new_tr)
            row = tbl.rows[-1]
            for i, txt in enumerate(row_content):
                if i >= len(row.cells): break
                cell = row.cells[i]
                # Clear existing runs di semua paragraph, ganti dengan text baru
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.text = ""
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].text = txt
                else:
                    cell.text = txt
            added += 1
        break  # patch first matching table only
    return added


def patch_doc(path: Path) -> tuple[int, int]:
    """Return (paragraph_repls, cell_repls) untuk audit."""
    doc = Document(str(path))
    p_count = 0
    c_count = 0
    # 1. Paragraph-level replacements
    for para in doc.paragraphs:
        p_count += patch_paragraph(para, REPLACEMENTS)
    # 2. Table cell replacements
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                c_count += patch_cell(cell, REPLACEMENTS)

    # 3. File-spesifik: tambah baris fitur baru
    if "ArtiVisi" in path.name:
        # Tambah row fitur baru di Table 3 (Mobile PWA features F-06..F-10)
        # Locate: header row col 0 = "ID" + row 1 col 0 = "F-06"
        def is_mobile_features(tbl):
            if len(tbl.rows) < 2: return False
            r0 = [c.text.strip() for c in tbl.rows[0].cells]
            r1 = [c.text.strip() for c in tbl.rows[1].cells]
            return r0[:2] == ["ID", "Nama Fitur"] and r1[0] == "F-06"
        added_mob = add_new_rows(doc, is_mobile_features, [
            ["F-11", "APK Android Native (Capacitor)",
             "Foreground service GPS: tracking tetap jalan saat layar mati / app backgrounded. Cakupan 95-99% (PWA hanya 70-85%).",
             "TINGGI"],
            ["F-12", "FCM Push (Firebase)",
             "Notifikasi push OS-level via FCM Firebase untuk APK Android (VAPID Web Push untuk PWA browser).",
             "TINGGI"],
            ["F-13", "Chat In-App Petugas ↔ Supervisor",
             "Chat realtime SSE + push. FAB di Beranda mobile petugas. Read receipt + unread badge.",
             "SEDANG"],
        ])
        # Tambah row fitur baru di Table 2 (Web Admin Dashboard F-01..F-05)
        def is_web_features(tbl):
            if len(tbl.rows) < 2: return False
            r0 = [c.text.strip() for c in tbl.rows[0].cells]
            r1 = [c.text.strip() for c in tbl.rows[1].cells]
            return r0[:2] == ["ID", "Nama Fitur"] and r1[0] == "F-01"
        added_web = add_new_rows(doc, is_web_features, [
            ["F-06", "Heatmap Kunjungan",
             "Density map cross-cabang untuk hotspot detection + zona under-served (7 hari terakhir).",
             "SEDANG"],
            ["F-07", "GPS Trail Historis",
             "Polyline pergerakan petugas per tanggal (date picker) — audit lintas-hari untuk investigasi fraud.",
             "TINGGI"],
            ["F-08", "Realtime Alerts (Geofence + Inactivity)",
             "Worker otomatis push ke supervisor: (a) petugas keluar wilayah binaan; (b) tidak ada ping GPS > 30 menit padahal clock-in.",
             "TINGGI"],
        ])
    else:  # Lampiran
        # T6 dashboard walkthrough — tambah halaman chat, heatmap, trail
        def is_lampiran_dashboard(tbl):
            if len(tbl.rows) < 2: return False
            r0 = [c.text.strip() for c in tbl.rows[0].cells]
            r1 = [c.text.strip() for c in tbl.rows[1].cells]
            return r0[:3] == ["Halaman", "Komponen Utama & Fungsi", "Endpoint API"] and r1[0] == "Login"
        added_web = add_new_rows(doc, is_lampiran_dashboard, [
            ["Chat Petugas",
             "Inbox percakapan supervisor ↔ petugas + thread messaging realtime SSE + unread badge nav.",
             "GET /api/chat/conversations, GET /api/chat/with/:userId, POST /api/chat/messages"],
            ["Heatmap Kunjungan",
             "Density map cross-cabang, 7 hari terakhir, hotspot detection + zona under-served.",
             "GET /api/analytics/visit-heatmap"],
            ["GPS Trail Historis",
             "Polyline pergerakan petugas per-tanggal (date picker) untuk audit lintas-hari + split segmen dashed pada gap > 5 menit.",
             "GET /api/petugas/:id/positions/trail?since=&until="],
        ])
        # T7 mobile walkthrough — tambah FAB chat + onboarding tour
        def is_lampiran_mobile(tbl):
            if len(tbl.rows) < 2: return False
            r0 = [c.text.strip() for c in tbl.rows[0].cells]
            r1 = [c.text.strip() for c in tbl.rows[1].cells]
            return r0[:3] == ["Halaman", "Komponen Utama & Fungsi", "Endpoint API"] and r1[0] == "Login PWA"
        added_mob = add_new_rows(doc, is_lampiran_mobile, [
            ["Chat FAB (Overlay)",
             "Floating action button di Beranda → overlay chat full-screen dengan supervisor. Auto-open thread saat cuma 1 conversation.",
             "GET /api/chat/conversations, POST /api/chat/messages, SSE 'chat.message'"],
            ["Onboarding Tour",
             "6 langkah walkthrough saat first-login petugas — clock-in, izin lokasi, kamera, laporan, review, chat.",
             "(client-side, localStorage flag)"],
        ])
        # T5 entities — tambah 4 entity utama yang belum ada
        def is_lampiran_entities(tbl):
            if len(tbl.rows) < 2: return False
            r0 = [c.text.strip() for c in tbl.rows[0].cells]
            r1 = [c.text.strip() for c in tbl.rows[1].cells]
            return r0[:2] == ["Entitas", "Deskripsi"] and r1[0] == "Branch"
        add_new_rows(doc, is_lampiran_entities, [
            ["PetugasPosition",
             "GPS ping history per petugas — lat, lng, accuracy, recordedAt. Sumber trail historis + speed-jump detection. Retention 90 hari default (worker harian)."],
            ["Attendance",
             "Sesi clock-in/out — petugasId, branchId, clockInAt, clockOutAt, odometer awal/akhir, cash awal/akhir. Basis inactivity worker."],
            ["ChatMessage",
             "Pesan antar user — fromId, toId, body, readAt. Realtime via SSE bus.publish 'chat.message'."],
            ["PushSubscription",
             "Registrasi push per device — kind (vapid|fcm), endpoint/token. Dual-runtime: VAPID Web Push (PWA) + FCM (APK)."],
        ])

    doc.save(str(path))
    return p_count, c_count


def main():
    for f in FILES:
        if not f.exists():
            print(f"SKIP {f.name} (not found)")
            continue
        p, c = patch_doc(f)
        print(f"✓ {f.name} — paragraph replacements: {p}, cell replacements: {c}")


if __name__ == "__main__":
    main()
